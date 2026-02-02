import sqlite3
import hashlib
import os
import uuid
import logging
from datetime import datetime
from faker import Faker
import openpyxl
from openpyxl import Workbook
from fpdf import FPDF
import PyPDF2
import json
import zipfile
import xml.etree.ElementTree as ET

logger = logging.getLogger(__name__)

class HoneyTokenManager:
    def __init__(self, db_path="honeytokens.db"):
        self.db_path = db_path
        self.fake = Faker()
        self.init_db()

    def init_db(self):
        """Инициализация базы данных с новой колонкой event_type"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Создаем таблицу токенов с колонкой event_type
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS tokens (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                token_guid TEXT UNIQUE NOT NULL,
                token_type TEXT NOT NULL,
                location TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                triggered INTEGER DEFAULT 0,
                triggered_at TIMESTAMP,
                ip_address TEXT,
                city TEXT,
                country TEXT,
                latitude REAL,
                longitude REAL,
                process_name TEXT,
                process_pid INTEGER,
                username TEXT,
                event_type TEXT DEFAULT 'unknown'  -- Новая колонка для типа события
            )
        ''')
        
        # Проверяем наличие колонки event_type (для обновления существующих БД)
        try:
            cursor.execute("PRAGMA table_info(tokens)")
            columns = [column[1] for column in cursor.fetchall()]
            if 'event_type' not in columns:
                cursor.execute("ALTER TABLE tokens ADD COLUMN event_type TEXT DEFAULT 'unknown'")
                logger.info("Добавлена колонка event_type в таблицу tokens")
        except Exception as e:
            logger.error(f"Ошибка проверки структуры таблицы: {e}")
        
        # Создаем таблицу пользователей для аутентификации
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                password_hash TEXT NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        ''')
        
        conn.commit()
        conn.close()


    def _calculate_file_hash(self, file_path):
        """Вычисление хеша файла"""
        try:
            hash_sha256 = hashlib.sha256()
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_sha256.update(chunk)
            return hash_sha256.hexdigest()
        except Exception as e:
            logger.error(f"Ошибка вычисления хеша файла {file_path}: {e}")
            return None

    def generate_file_token(self, file_path, content=None, use_faker=True, obfuscate_guid=True):
        """Генерация файла-ловушки с поддержкой разных форматов"""
        token_guid = str(uuid.uuid4())
        file_ext = os.path.splitext(file_path)[1].lower()

        # Создаем директорию, если она не существует
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        if file_ext == '.pdf':
            self._generate_pdf_token(file_path, token_guid, use_faker, obfuscate_guid)
        elif file_ext in ['.xlsx', '.xls']:
            self._generate_excel_token(file_path, token_guid, use_faker, obfuscate_guid)
        elif file_ext == '.docx':
            self._generate_word_token(file_path, token_guid, use_faker, obfuscate_guid)
        else:
            self._generate_text_token(file_path, token_guid, content, use_faker, obfuscate_guid)

        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO tokens (token_guid, token_type, location)
            VALUES (?, ?, ?)
        ''', (token_guid, 'file', file_path))
        conn.commit()
        conn.close()

        logger.info(f"Файл-ловушка создан: {file_path}, GUID: {token_guid}")
        print(f"✅ Файл-ловушка создан: {file_path}")
        print(f"   GUID токена: {token_guid}")
        return token_guid

    def _generate_text_token(self, file_path, token_guid, content, use_faker, obfuscate_guid):
        """Генерация текстового файла-ловушки"""
        if use_faker:
            content = f"""CONFIDENTIAL

Company employee data:
Name: {self.fake.name()}
Email: {self.fake.email()}
Phone: {self.fake.phone_number()}
Address: {self.fake.address()}

System credentials:
Username: admin_{self.fake.user_name()}
Password: {self.fake.password(length=12)}

Database connection:
Host: db-{self.fake.word()}.company.com
Port: {self.fake.random_int(min=3000, max=9999)}

Access restricted to authorized personnel only.
"""
        
        if obfuscate_guid:
            hidden_guid = f"\n\n<!-- DocumentID: {token_guid} -->\nRevision: 1.0"
            content += hidden_guid
        else:
            content += f"\n\nToken identifier: {token_guid}\nCreated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(content)

    def _generate_pdf_token(self, file_path, token_guid, use_faker, obfuscate_guid):
        """Генерация PDF файла-ловушки с поддержкой Unicode"""
        try:
            # Создаем PDF с поддержкой Unicode
            pdf = FPDF()
            
            # Добавляем шрифт с поддержкой Unicode (DejaVuSans)
            # Сначала попробуем использовать стандартные шрифты
            pdf.add_page()
            
            if use_faker:
                # Используем английский текст для PDF
                content = [
                    "CONFIDENTIAL DOCUMENT",
                    f"Employee Report - {datetime.now().strftime('%d.%m.%Y')}",
                    "",
                    f"Manager: {self.fake.name()}",
                    f"Department: {self.fake.job()}",
                    f"Report Date: {self.fake.date()}",
                    "",
                    "Employee List:"
                ]
                
                for i in range(5):
                    content.append(f"{i+1}. {self.fake.name()} - {self.fake.email()}")
            else:
                content = ["This is a confidential PDF honeytoken document."]

            # Устанавливаем шрифт
            pdf.set_font("Arial", size=12)
            
            for line in content:
                try:
                    pdf.cell(200, 10, txt=line, ln=True)
                except:
                    # Если есть проблемы с кодировкой, используем альтернативный текст
                    pdf.cell(200, 10, txt="Confidential Document - Access Restricted", ln=True)
        
            if obfuscate_guid:
                # Добавляем GUID в метаданные
                pdf.set_title(f"Report_{datetime.now().strftime('%Y%m%d')}")
                pdf.set_author("HR Department")
                pdf.set_subject(f"TokenID: {token_guid}")
            else:
                pdf.cell(200, 10, txt=f"TokenID: {token_guid}", ln=True)

            pdf.output(file_path)
            
        except Exception as e:
            logger.error(f"Ошибка генерации PDF: {e}")
            # Создаем простой текстовый файл вместо PDF
            logger.info(f"Создаем текстовый файл вместо PDF: {file_path}")
            self._generate_text_token(file_path.replace('.pdf', '.txt'), token_guid, None, use_faker, obfuscate_guid)

    def _generate_excel_token(self, file_path, token_guid, use_faker, obfuscate_guid):
        """Генерация Excel файла-ловушки"""
        wb = Workbook()
        ws = wb.active
        ws.title = "Employees"

        # Заголовки на английском
        headers = ["ID", "Full Name", "Position", "Email", "Phone", "Department"]
        ws.append(headers)

        # Данные
        if use_faker:
            for i in range(10):
                row = [
                    i + 1,
                    self.fake.name(),
                    self.fake.job(),
                    self.fake.email(),
                    self.fake.phone_number(),
                    self.fake.word().capitalize()
                ]
                ws.append(row)
        else:
            ws.append([1, "Test Employee", "Manager", "test@company.com", "+79990000000", "IT"])

        if obfuscate_guid:
            # Скрываем GUID в скрытой ячейке
            ws['Z100'] = token_guid
            ws['Z100'].font = openpyxl.styles.Font(color="FFFFFF")  # Белый текст
        else:
            ws.append(["", f"TokenID: {token_guid}"])

        wb.save(file_path)

    def _generate_word_token(self, file_path, token_guid, use_faker, obfuscate_guid):
        """Генерация реального Word документа-ловушки"""
        try:
            from docx import Document
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Добавляем заголовок
            title = doc.add_heading('CONFIDENTIAL BUSINESS PLAN', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Основная информация
            doc.add_paragraph(f"Company: {self.fake.company()}")
            doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(f"Author: {self.fake.name()}")
            doc.add_paragraph(f"Document ID: BP-{self.fake.random_number(digits=6)}")
            doc.add_paragraph("")
            
            # Executive Summary
            doc.add_heading('EXECUTIVE SUMMARY', level=1)
            doc.add_paragraph(self.fake.paragraph(nb_sentences=5))
            doc.add_paragraph("")
            
            # Market Analysis
            doc.add_heading('MARKET ANALYSIS', level=1)
            market_content = [
                f"• Target Market: {self.fake.catch_phrase()}",
                f"• Market Size: ${self.fake.random_number(digits=3)} million",
                f"• Growth Rate: {self.fake.random_int(min=5, max=25)}% annually",
                f"• Key Competitors: {self.fake.company()}, {self.fake.company()}"
            ]
            for item in market_content:
                doc.add_paragraph(item)
            doc.add_paragraph(self.fake.paragraph(nb_sentences=3))
            doc.add_paragraph("")
            
            # Financial Projections
            doc.add_heading('FINANCIAL PROJECTIONS', level=1)
            financial_data = [
                f"Year 1 Revenue: ${self.fake.random_number(digits=7):,}",
                f"Year 2 Revenue: ${self.fake.random_number(digits=7):,}",
                f"Year 3 Revenue: ${self.fake.random_number(digits=8):,}",
                f"Projected Profit Margin: {self.fake.random_int(min=15, max=40)}%"
            ]
            for item in financial_data:
                doc.add_paragraph(item)
            doc.add_paragraph("")
            
            # Growth Strategy
            doc.add_heading('GROWTH STRATEGY', level=1)
            doc.add_paragraph(self.fake.paragraph(nb_sentences=6))
            doc.add_paragraph("")
            
            # Risk Assessment
            doc.add_heading('RISK ASSESSMENT', level=1)
            risks = [
                f"• Market Risk: {self.fake.sentence()}",
                f"• Operational Risk: {self.fake.sentence()}",
                f"• Financial Risk: {self.fake.sentence()}"
            ]
            for risk in risks:
                doc.add_paragraph(risk)
            doc.add_paragraph("")
            
            # Confidential Notice
            confidential = doc.add_paragraph()
            confidential.alignment = WD_ALIGN_PARAGRAPH.CENTER
            confidential.add_run("CONFIDENTIAL AND PROPRIETARY").bold = True
            
            notice = doc.add_paragraph()
            notice.alignment = WD_ALIGN_PARAGRAPH.CENTER
            notice.add_run("This document contains trade secrets and confidential information of ")
            notice.add_run(f"{self.fake.company()}").bold = True
            notice.add_run(".")
            
            warning = doc.add_paragraph()
            warning.alignment = WD_ALIGN_PARAGRAPH.CENTER
            warning.add_run("Unauthorized access, copying, or distribution is strictly prohibited.").bold = True
            
            # Добавляем GUID в документ (скрыто или явно)
            if not obfuscate_guid:
                doc.add_paragraph(f"Document Token: {token_guid}")
            else:
                # Добавляем скрытый GUID в метаданные
                self._add_guid_to_word_metadata(file_path, token_guid, doc)
            
            # Сохраняем документ
            doc.save(file_path)
            logger.info(f"Создан Word документ: {file_path}")
                
        except Exception as e:
            logger.error(f"Ошибка генерации Word документа: {e}")
            # Fallback - создаем текстовый файл
            self._generate_text_token(file_path, token_guid, None, use_faker, obfuscate_guid)

    def _add_guid_to_word_metadata(self, file_path, token_guid, doc):
        """Добавляет GUID в метаданные Word документа"""
        try:
            # Сначала сохраняем документ
            doc.save(file_path)
            
            # Затем открываем как zip и модифицируем метаданные
            with zipfile.ZipFile(file_path, 'a') as docx_zip:
                # Читаем core.xml
                core_xml = docx_zip.read('docProps/core.xml')
                root = ET.fromstring(core_xml)
                
                # Находим или создаем поле subject
                subject_elem = None
                for elem in root:
                    if 'subject' in elem.tag:
                        subject_elem = elem
                        break
                
                if subject_elem is None:
                    # Создаем новый элемент subject
                    ns = {'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties'}
                    subject_elem = ET.Element('{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}subject')
                    root.append(subject_elem)
                
                subject_elem.text = f"Business Plan - TokenID: {token_guid}"
                
                # Сохраняем обратно
                docx_zip.writestr('docProps/core.xml', ET.tostring(root, encoding='unicode'))
                
        except Exception as e:
            logger.warning(f"Не удалось добавить GUID в метаданные Word: {e}")
            # Добавляем GUID как скрытый текст в конец документа
            doc.add_paragraph(f"Reviewer: {self.fake.name()} - ID: {token_guid}")

    def check_token_triggered(self, token_guid):
        """Проверка, сработал ли токен"""
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        cursor.execute('SELECT triggered FROM tokens WHERE token_guid = ?', (token_guid,))
        result = cursor.fetchone()
        conn.close()
        
        # Возвращаем 0 если запись не найдена или triggered=0
        if result is None:
            logger.debug(f"Токен {token_guid} не найден в БД")
            return 0
        return result[0] if result else 0

    def mark_token_triggered(self, token_guid, ip=None, process_info=None, event_type='modify'):
        """Пометить токен как сработавший с указанием типа события"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            logger.debug(f"Обновление токена {token_guid}: event_type={event_type}, ip={ip}")
            
            update_data = [1, datetime.now(), event_type]  # Добавлен event_type
            update_fields = "triggered = ?, triggered_at = ?, event_type = ?"
            
            if ip:
                update_fields += ", ip_address = ?"
                update_data.append(ip)
                logger.info(f"Установлен IP для токена {token_guid}: {ip}")
            else:
                logger.warning(f"IP адрес не определен для токена {token_guid}")
            
            if process_info:
                update_fields += ", process_name = ?, process_pid = ?, username = ?"
                update_data.extend([
                    process_info.get('name', 'Unknown'),
                    process_info.get('pid', 0),
                    process_info.get('username', 'Unknown')
                ])
        
            update_data.append(token_guid)
            
            cursor.execute(f'''
                UPDATE tokens 
                SET {update_fields}
                WHERE token_guid = ?
            ''', update_data)
            
            rows_updated = cursor.rowcount
            conn.commit()
            
            if rows_updated > 0:
                logger.warning(f"Honey token {token_guid} был активирован! IP: {ip}, Тип: {event_type}, Обновлено строк: {rows_updated}")
                print(f"🚨 ТРЕВОГА: Honey token {token_guid} был активирован! IP: {ip}, Тип: {event_type}")
            else:
                logger.error(f"Токен {token_guid} не найден при обновлении!")
                print(f"❌ ОШИБКА: Токен {token_guid} не найден в БД!")
            
        except sqlite3.Error as e:
            logger.error(f"Ошибка БД при обновлении токена {token_guid}: {e}")
        finally:
            if conn:
                conn.close()

    def update_token_geo(self, token_guid, geo_data):
        """Обновление геоданных токена"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute('''
                UPDATE tokens 
                SET city = ?, country = ?, latitude = ?, longitude = ?
                WHERE token_guid = ?
            ''', (
                geo_data.get('city'), 
                geo_data.get('country'), 
                geo_data.get('lat'), 
                geo_data.get('lng'), 
                token_guid
            ))
            conn.commit()
            logger.info(f"Геоданные обновлены для токена {token_guid}")
        except sqlite3.Error as e:
            logger.error(f"Ошибка БД при обновлении геоданных для токена {token_guid}: {e}")
        finally:
            if conn:
                conn.close()

    def get_active_file_tokens(self):
        """Получить все активные (нетриггернутые) файловые токены"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT token_guid, location FROM tokens WHERE token_type = 'file' AND triggered = 0")
            tokens = cursor.fetchall()
            return tokens
        except Exception as e:
            logger.error(f"Ошибка получения активных файловых токенов: {e}")
            return []
        finally:
            if conn:
                conn.close()

    def get_all_tokens(self):
        """Получить все токены (для отладки)"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM tokens")
            tokens = cursor.fetchall()
            return tokens
        except Exception as e:
            logger.error(f"Ошибка получения всех токенов: {e}")
            return []
        finally:
            if conn:
                conn.close()

    def get_triggered_tokens_today(self):
        """Получить токены, сработавшие сегодня (безопасная версия)"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM tokens 
                WHERE triggered = 1 AND DATE(triggered_at) = DATE('now')
            ''')
            tokens = cursor.fetchall()
            return tokens
        except Exception as e:
            logger.error(f"Ошибка получения сегодняшних токенов: {e}")
            return []
        finally:
            if conn:
                conn.close()

    def generate_trap_tokens(self, original_path, levels=2):
        """Генерация дополнительных ловушек после срабатывания (максимум levels)"""
        try:
            # Проверяем, существует ли оригинальный файл
            if not os.path.exists(original_path):
                logger.warning(f"Оригинальный файл не найден: {original_path}")
                return []
            
            # Проверяем, не является ли файл временным
            if self._is_temporary_file(original_path):
                logger.warning(f"Игнорируем временный файл для создания ловушек: {original_path}")
                return []
            
            base_name = os.path.splitext(os.path.basename(original_path))[0]
            ext = os.path.splitext(original_path)[1].lower()
            directory = os.path.dirname(original_path)
            
            trap_tokens = []
            
            # Проверяем, поддерживается ли формат файла для создания ловушек
            supported_extensions = ['.txt', '.pdf', '.docx', '.xlsx', '.xls']
            if ext not in supported_extensions:
                logger.warning(f"Формат файла {ext} не поддерживается для создания ловушек")
                return []
            
            # Подсчитываем существующие ловушки
            existing_traps = self.count_existing_traps(original_path)
            if existing_traps >= levels:
                logger.info(f"Для файла {original_path} уже создано {existing_traps} ловушек, пропускаем создание новых")
                return []
            
            traps_to_create = min(levels - existing_traps, levels)
            
            for i in range(traps_to_create):
                # Создаем имя файла-ловушки
                trap_name = f"backup_{base_name}_v{i+1+existing_traps}{ext}"
                trap_path = os.path.join(directory, trap_name)
                
                # Проверяем, не существует ли уже такой файл
                if os.path.exists(trap_path):
                    logger.debug(f"Ловушка уже существует: {trap_path}")
                    continue
                
                # Генерируем новый токен
                trap_guid = self.generate_file_token(
                    trap_path,
                    use_faker=True,
                    obfuscate_guid=True
                )
                trap_tokens.append((trap_guid, trap_path))
                
                logger.info(f"Создана ловушка: {trap_path}, GUID: {trap_guid}")
            
            logger.info(f"Сгенерировано {len(trap_tokens)} ловушек для {original_path}")
            return trap_tokens
            
        except Exception as e:
            logger.error(f"Ошибка генерации ловушек для {original_path}: {e}")
            return []
    
    def count_existing_traps(self, original_file_path):
        """Подсчет существующих ловушек для оригинального файла"""
        try:
            base_name = os.path.splitext(os.path.basename(original_file_path))[0]
            ext = os.path.splitext(original_file_path)[1].lower()
            directory = os.path.dirname(original_file_path)
            
            if not os.path.exists(directory):
                return 0
            
            trap_count = 0
            for filename in os.listdir(directory):
                if filename.startswith(f"backup_{base_name}_v") and filename.endswith(ext):
                    trap_count += 1
            
            return trap_count
        except Exception as e:
            logger.error(f"Ошибка подсчета ловушек: {e}")
            return 0
    
    def _is_temporary_file(self, file_path):
        """Проверка, является ли файл временным"""
        filename = os.path.basename(file_path)
        
        # Игнорируем временные файлы Office
        if filename.startswith('~$') or filename.startswith('.~'):
            return True
        
        # Игнорируем временные файлы с расширением .tmp
        if filename.endswith('.tmp'):
            return True
        
        # Игнорируем системные файлы
        if filename.startswith('~') or filename.startswith('._'):
            return True
        
        return False

    def safe_get_token_field(self, token, index, default=None):
        """Безопасное получение поля токена по индексу"""
        try:
            if token and len(token) > index:
                value = token[index]
                return value if value is not None else default
            return default
        except (IndexError, TypeError, AttributeError) as e:
            logger.debug(f"Ошибка получения поля {index} токена: {e}")
            return default

    def get_database_structure(self):
        """Получить структуру базы данных для отладки"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Получаем структуру таблицы
            cursor.execute("PRAGMA table_info(tokens)")
            columns = cursor.fetchall()
            
            # Получаем статистику
            cursor.execute("SELECT COUNT(*) FROM tokens")
            total = cursor.fetchone()[0]
            cursor.execute("SELECT COUNT(*) FROM tokens WHERE triggered = 1")
            triggered = cursor.fetchone()[0]
            cursor.execute("SELECT COUNT(*) FROM tokens WHERE token_type = 'file'")
            files = cursor.fetchone()[0]
            
            # Получаем информацию о сработавших токенах с IP
            cursor.execute("SELECT COUNT(*) FROM tokens WHERE triggered = 1 AND ip_address IS NOT NULL")
            with_ip = cursor.fetchone()[0]
            
            return {
                'columns': columns,
                'stats': {
                    'total': total,
                    'triggered': triggered,
                    'active': total - triggered,
                    'files': files,
                    'triggered_with_ip': with_ip
                }
            }
        except Exception as e:
            logger.error(f"Ошибка получения структуры БД: {e}")
            return None
        finally:
            if conn:
                conn.close()

    def cleanup_old_tokens(self, days=30):
        """Очистка старых сработавших токенов"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute('''
                DELETE FROM tokens 
                WHERE triggered = 1 AND triggered_at < DATE('now', ?)
            ''', (f'-{days} days',))
            deleted_count = cursor.rowcount
            conn.commit()
            logger.info(f"Удалено {deleted_count} старых токенов (старше {days} дней)")
            return deleted_count
        except Exception as e:
            logger.error(f"Ошибка очистки старых токенов: {e}")
            return 0
        finally:
            if conn:
                conn.close()

    def get_token_by_guid(self, token_guid):
        """Получить токен по GUID"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM tokens WHERE token_guid = ?", (token_guid,))
            token = cursor.fetchone()
            return token
        except Exception as e:
            logger.error(f"Ошибка получения токена {token_guid}: {e}")
            return None
        finally:
            if conn:
                conn.close()

    def get_triggered_tokens_with_geo(self):
        """Получить сработавшие токены с геоданными"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute('''
                SELECT * FROM tokens 
                WHERE triggered = 1 AND latitude IS NOT NULL AND longitude IS NOT NULL
            ''')
            tokens = cursor.fetchall()
            return tokens
        except Exception as e:
            logger.error(f"Ошибка получения токенов с геоданными: {e}")
            return []
        finally:
            if conn:
                conn.close()

    def get_token_by_file_path(self, file_path):
        """Получить токен по пути к файлу"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            # Используем точное соответствие пути
            cursor.execute("SELECT * FROM tokens WHERE location = ?", (file_path,))
            token = cursor.fetchone()
            
            return token
        except Exception as e:
            logger.error(f"Ошибка получения токена по пути {file_path}: {e}")
            return None
        finally:
            if conn:
                conn.close()

    def get_tokens_by_folder(self, folder_path):
        """Получить токены по папке"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            cursor.execute("SELECT * FROM tokens WHERE location LIKE ?", (f"{folder_path}%",))
            tokens = cursor.fetchall()
            return tokens
        except Exception as e:
            logger.error(f"Ошибка получения токенов по папке {folder_path}: {e}")
            return []
        finally:
            if conn:
                conn.close()            

    def export_tokens_to_json(self, file_path="tokens_export.json"):
        """Экспорт всех токенов в JSON файл"""
        try:
            tokens = self.get_all_tokens()
            tokens_data = []
            
            for token in tokens:
                token_data = {
                    'id': self.safe_get_token_field(token, 0),
                    'guid': self.safe_get_token_field(token, 1),
                    'type': self.safe_get_token_field(token, 2),
                    'location': self.safe_get_token_field(token, 3),
                    'created_at': self.safe_get_token_field(token, 4),
                    'triggered': bool(self.safe_get_token_field(token, 5, 0)),
                    'triggered_at': self.safe_get_token_field(token, 6),
                    'ip_address': self.safe_get_token_field(token, 7),
                    'city': self.safe_get_token_field(token, 8),
                    'country': self.safe_get_token_field(token, 9),
                    'latitude': self.safe_get_token_field(token, 10),
                    'longitude': self.safe_get_token_field(token, 11),
                    'process_name': self.safe_get_token_field(token, 12),
                    'process_pid': self.safe_get_token_field(token, 13),
                    'username': self.safe_get_token_field(token, 14)
                }
                tokens_data.append(token_data)
            
            with open(file_path, 'w', encoding='utf-8') as f:
                json.dump(tokens_data, f, indent=2, ensure_ascii=False, default=str)
            
            logger.info(f"Токены экспортированы в {file_path}")
            return True
        except Exception as e:
            logger.error(f"Ошибка экспорта токенов: {e}")
            return False

    def get_token_statistics(self):
        """Получить расширенную статистику по токенам"""
        try:
            conn = sqlite3.connect(self.db_path)
            cursor = conn.cursor()
            
            stats = {}
            
            # Основная статистика
            cursor.execute("SELECT COUNT(*) FROM tokens")
            stats['total'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM tokens WHERE triggered = 1")
            stats['triggered'] = cursor.fetchone()[0]
            
            cursor.execute("SELECT COUNT(*) FROM tokens WHERE triggered = 0")
            stats['active'] = cursor.fetchone()[0]
            
            # Статистика по типам
            cursor.execute("SELECT token_type, COUNT(*) FROM tokens GROUP BY token_type")
            stats['by_type'] = dict(cursor.fetchall())
            
            # Статистика по дням
            cursor.execute('''
                SELECT DATE(created_at), COUNT(*) 
                FROM tokens 
                GROUP BY DATE(created_at) 
                ORDER BY DATE(created_at) DESC 
                LIMIT 7
            ''')
            stats['created_last_7_days'] = dict(cursor.fetchall())
            
            # Топ IP адресов
            cursor.execute('''
                SELECT ip_address, COUNT(*) 
                FROM tokens 
                WHERE ip_address IS NOT NULL 
                GROUP BY ip_address 
                ORDER BY COUNT(*) DESC 
                LIMIT 10
            ''')
            stats['top_ips'] = dict(cursor.fetchall())
            
            # Гео статистика
            cursor.execute("SELECT COUNT(*) FROM tokens WHERE city IS NOT NULL")
            stats['with_geo'] = cursor.fetchone()[0]
            
            return stats
            
        except Exception as e:
            logger.error(f"Ошибка получения статистики: {e}")
            return {}
        finally:
            if conn:
                conn.close()

# Вспомогательные функции
def setup_database_backup(db_path, backup_dir="backups"):
    """Настройка автоматического бэкапа базы данных"""
    try:
        os.makedirs(backup_dir, exist_ok=True)
        backup_file = os.path.join(backup_dir, f"honeytokens_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.db")
        
        import shutil
        shutil.copy2(db_path, backup_file)
        logger.info(f"Создан бэкап базы данных: {backup_file}")
        return backup_file
    except Exception as e:
        logger.error(f"Ошибка создания бэкапа: {e}")
        return None
def delete_folder_tokens(self, folder_path):
    """Удаление всех токенов в указанной папке"""
    try:
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Удаляем токены по пути папки
        cursor.execute("DELETE FROM tokens WHERE location LIKE ?", (f"{folder_path}%",))
        deleted_count = cursor.rowcount
        conn.commit()
        
        logger.info(f"Удалено {deleted_count} токенов из папки {folder_path}")
        return deleted_count
        
    except Exception as e:
        logger.error(f"Ошибка удаления токенов папки {folder_path}: {e}")
        return 0
    finally:
        if conn:
            conn.close()

def get_tokens_in_folder(self, folder_path):
    """Получить все токены в указанной папке"""
    try:
        conn = sqlite3.connect(self.db_path)
        cursor = conn.cursor()
        
        # Получаем токены по пути папки
        cursor.execute("SELECT * FROM tokens WHERE location LIKE ?", (f"{folder_path}%",))
        tokens = cursor.fetchall()
        
        return tokens
        
    except Exception as e:
        logger.error(f"Ошибка получения токенов из папки {folder_path}: {e}")
        return []
    finally:
        if conn:
            conn.close()
# Пример использования
if __name__ == "__main__":
    import json
    
    # Настройка логирования
    logging.basicConfig(
        level=logging.INFO,
        format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
    )
    
    manager = HoneyTokenManager()
    
    # Проверка структуры БД
    structure = manager.get_database_structure()
    if structure:
        print("📊 Структура базы данных:")
        for col in structure['columns']:
            print(f"  {col[1]} ({col[2]})")
        print(f"\n📈 Статистика:")
        print(f"  Всего токенов: {structure['stats']['total']}")
        print(f"  Активных: {structure['stats']['active']}")
        print(f"  Сработавших: {structure['stats']['triggered']}")
        print(f"  Файловых: {structure['stats']['files']}")
        print(f"  Сработавших с IP: {structure['stats']['triggered_with_ip']}")
    
    # Создание тестовых токенов
    print("\n🧪 Создание тестовых токенов...")
    
    # Файловый токен
    test_file_token = manager.generate_file_token("test_document.pdf")
    print(f"📄 Создан файловый токен: {test_file_token}")
    
    # Текстовый токен
    test_text_token = manager.generate_file_token("test_secret.txt")
    print(f"📝 Создан текстовый токен: {test_text_token}")
    
    # Word токен
    test_word_token = manager.generate_file_token("test_business_plan.docx")
    print(f"📝 Создан Word токен: {test_word_token}")
    
    # Расширенная статистика
    stats = manager.get_token_statistics()
    print(f"\n📊 Расширенная статистика:")
    print(json.dumps(stats, indent=2, ensure_ascii=False, default=str))
    
    # Создание бэкапа
    backup_file = setup_database_backup("honeytokens.db")
    if backup_file:
        print(f"💾 Создан бэкап: {backup_file}")